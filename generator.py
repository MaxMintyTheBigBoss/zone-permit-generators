# -*- coding: utf-8 -*-
"""
Генерация документов пропусков по п. 14.3.

Использует общие модули permit_common (пути, замена плейсхолдеров, join)
и permit_db (SQLite-хранилище истории). API внизу сохранён совместимым
с прежним JSON-слоем: _db_load/_db_remember/db_find.
"""
import os
import sys
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

from permit_common import (
    resource_path, runtime_base, template_dir,
    load_reference, filter_objects_by_districts,
    join_districts, join_objects, sanitize, today_dmy,
    fill_doc, _inline_replace, _PH_RE,
)
import permit_db as dbm

# Цель въезда (Placeholder_6).
GOAL_BLAGOUSTROISTVO = "благоустройство места захоронения"
GOAL_CUSTOM = "свой вариант"

# Районы (Placeholder_4) — множественный выбор из 8.
DISTRICTS = [
    "Брагинский", "Буда-Кошелевский", "Ветковский", "Добрушский",
    "Кормянский", "Наровлянский", "Хойникский", "Чечерский",
]

# Кому на подписание (Placeholder_20) — полный список.
SIGNERS = [
    "Заместитель начальника отдела Путькова Т.М.",
    "Начальник отдела Соломейчук А.В.",
    "Заместитель начальника главного управления В.О.Шабловский",
    "Главный специалист Гвоздарев А.А.",
    "Главный специалист Геращенко Г.Н.",
    "Главный специалист Колесан А.И.",
    "Главный специалист Курило А.В.",
    "Главный специалист Новик П.Н.",
    "Главный специалист Одиноченко И.В.",
    "Главный специалист Першко А.С.",
]


def _iter_all_paragraphs(doc):
    """Все абзацы документа, включая ячейки таблиц."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _set_para_text(p, text):
    """Записывает текст в абзац, сохраняя форматирование первого run."""
    runs = p.findall(qn('w:r'))
    if not runs:
        return
    r0 = runs[0]
    for t in r0.findall(qn('w:t')):
        r0.remove(t)
    t = r0.makeelement(qn('w:t'), {})
    t.text = str(text)
    r0.append(t)
    for r in runs[1:]:
        for t in r.findall(qn('w:t')):
            t.text = ""


def _build_persons_block(doc, persons):
    """Перестраивает блок сопровождающих в заявлении (по одному на человека)."""
    W = qn('w:p')
    WTR = qn('w:tr')
    WTC = qn('w:tc')
    TBL = qn('w:tbl')
    anchor_text = "для граждан Республики Беларусь"

    tbl = None
    for t in doc.element.body.iter(TBL):
        full = "".join(x.text or "" for x in t.iter(qn('w:t')))
        if "Placeholder_10.1" in full:
            tbl = t
            break
    if tbl is None:
        return
    rows = [tr for tr in tbl if tr.tag == WTR]

    tpl_row = None
    anchor_row = None
    for tr in rows:
        txt = "".join(c.text for c in tr.iter(qn('w:t')))
        if tpl_row is None and "Placeholder_10.1" in txt:
            tpl_row = tr
        if anchor_row is None and anchor_text in txt:
            anchor_row = tr
        if tpl_row is not None and anchor_row is not None:
            break
    if tpl_row is None or anchor_row is None:
        return

    num_tpl = deepcopy(tpl_row)
    data_tpl = deepcopy(tpl_row)

    try:
        i = rows.index(tpl_row)
        for tr in rows[i:i + 2]:
            if tr is not None and tr in tbl:
                tbl.remove(tr)
    except Exception:
        pass

    for i, pers in enumerate(persons):
        name = " ".join(x for x in [pers.get("last_name", ""), pers.get("first_name", ""), pers.get("middle_name", "")] if x)
        birth = pers.get("birth_date", "")
        line = (name + (", " if name and birth else "") + birth).strip()
        new_row = deepcopy(data_tpl)
        cells = [c for c in new_row if c.tag == WTC]
        if len(cells) >= 2:
            k1 = [p for p in cells[1] if p.tag == W]
            if k1:
                _set_para_text(k1[0], line)
                for pp in k1[1:]:
                    for rr in pp.findall(qn('w:r')):
                        for tt in rr.findall(qn('w:t')):
                            tt.text = ""
            k0 = [p for p in cells[0] if p.tag == W]
            if k0:
                _set_para_text(k0[0], str(i + 1))
                for pp in k0[1:]:
                    for rr in pp.findall(qn('w:r')):
                        for tt in rr.findall(qn('w:t')):
                            tt.text = ""
        anchor_row.addprevious(new_row)


def _today_dmy():
    return today_dmy()


def make_application(data):
    """Заявление 14.3."""
    doc = Document(os.path.join(template_dir(), "Заявление_14.3.docx"))
    districts = data.get("districts", [])
    m = {
        "Placeholder_1.1": data["last_name"],
        "Placeholder_1.2": data["first_name"],
        "Placeholder_1.3": data["middle_name"],
        "Placeholder_2": data.get("birth_date", ""),
        "Placeholder_3": data.get("id_number", ""),
        "Placeholder_4": join_districts(districts),
        "Placeholder_5": data.get("objects", ""),
        "Placeholder_6": data.get("goal", ""),
        "Placeholder_7": data.get("date_from", ""),
        "Placeholder_8": data.get("date_to", ""),
        "Placeholder_12": data.get("car_make", ""),
        "Placeholder_13": data.get("car_number", ""),
        "Placeholder_9": data.get("app_date") or _today_dmy(),
    }
    for i in range(3):
        m["Placeholder_4.%d" % (i + 1)] = districts[i] if i < len(districts) else ""

    fill_doc(doc, m)
    _build_persons_block(doc, data.get("persons", []))
    return doc


def make_individual(data, person=None):
    """Индивидуальный пропуск. person = None => заявитель, иначе пассажир."""
    person = person if person is not None else data
    doc = Document(os.path.join(template_dir(), "Пропуск_индивидуальный.docx"))
    m = {
        "Placeholder_1.1": person["last_name"],
        "Placeholder_1.2": person["first_name"],
        "Placeholder_1.3": person["middle_name"],
        "Placeholder_4": join_districts(data.get("districts", [])),
        "Placeholder_5": data.get("objects", ""),
        "Placeholder_6": data.get("goal", ""),
        "Placeholder_7": data.get("date_from", ""),
        "Placeholder_8": data.get("date_to", ""),
        "Placeholder_20": data.get("issued_by", ""),
    }
    fill_doc(doc, m)
    return doc


def make_transport(data):
    """Транспортный пропуск (заинтересованное лицо — заявитель)."""
    doc = Document(os.path.join(template_dir(), "Пропуск_транспортный.docx"))
    m = {
        "Placeholder_1.1": data["last_name"],
        "Placeholder_1.2": data["first_name"],
        "Placeholder_1.3": data["middle_name"],
        "Placeholder_4": join_districts(data.get("districts", [])),
        "Placeholder_5": data.get("objects", ""),
        "Placeholder_6": data.get("goal", ""),
        "Placeholder_7": data.get("date_from", ""),
        "Placeholder_8": data.get("date_to", ""),
        "Placeholder_12": data.get("car_make", ""),
        "Placeholder_13": data.get("car_number", ""),
        "Placeholder_20": data.get("issued_by", ""),
    }
    fill_doc(doc, m)
    return doc


def generate_all(data, output_dir=None):
    """Полный комплект документов по п. 14.3. Возвращает список файлов."""
    if output_dir is None:
        output_dir = os.path.join(runtime_base(), "output")
    stamp = datetime.now().strftime("%d.%m.%Y.%H.%M")
    out = os.path.join(output_dir, stamp)
    os.makedirs(out, exist_ok=True)

    created = []
    lname = sanitize(data["last_name"], "Заявитель")

    d = make_application(data)
    p = os.path.join(out, "Заявление_%s.docx" % lname)
    d.save(p)
    created.append(p)

    d = make_individual(data)
    p = os.path.join(out, "Пропуск_Заявитель_%s.docx" % lname)
    d.save(p)
    created.append(p)

    for person in data.get("persons", []):
        plname = sanitize(person["last_name"], "Пассажир")
        d = make_individual(data, person)
        p = os.path.join(out, "Пропуск_Пассажир_%s.docx" % plname)
        d.save(p)
        created.append(p)

    if data.get("car_make", "").strip() or data.get("car_number", "").strip():
        car = sanitize(data.get("car_number", "") or "Авто", "Авто")
        d = make_transport(data)
        p = os.path.join(out, "Пропуск_Транспорт_%s.docx" % car)
        d.save(p)
        created.append(p)

    _db_remember(data)
    return created


# ---------------------------------------------------------------------------
# База знаний — SQLite (совместимый интерфейс с прежним JSON-слоем)
# ---------------------------------------------------------------------------
def _db_path():
    d = os.path.join(runtime_base(), "data")
    os.makedirs(d, exist_ok=True)
    return os.path.join(d, dbm.db_filename("143"))


def _db_conn():
    return dbm.PermitDB(_db_path())


def _db_load():
    """Все записи (для выпадающих списков)."""
    try:
        return _db_conn().load()
    except Exception:
        return []


def _db_load_fio():
    try:
        return _db_conn().load("fio")
    except Exception:
        return []


def db_find(fio):
    """Последняя запись по ФИО."""
    return dbm.PermitDB(_db_path()).find(fio, "fio") if (fio or "").strip() else None


def _db_remember(data):
    """Сохраняет запись (с миграцией JSON при первом запуске)."""
    db = _db_conn()
    # однократная миграция из прежнего JSON
    try:
        import json as _json
        legacy = os.path.join(runtime_base(), "data", "permit_history.json")
        if os.path.exists(legacy):
            recs = _json.load(open(legacy, encoding="utf-8"))
            if recs:
                dbm.migrate_from_json(db, legacy, "fio")
                os.rename(legacy, legacy + ".migrated")
    except Exception:
        pass
    key = " ".join(x for x in [data.get("last_name", ""), data.get("first_name", ""), data.get("middle_name", "")] if x).strip()
    rec = {
        "key": key, "last_name": data.get("last_name", ""), "first_name": data.get("first_name", ""),
        "middle_name": data.get("middle_name", ""), "birth_date": data.get("birth_date", ""),
        "id_number": data.get("id_number", ""), "goal": data.get("goal", ""),
        "districts": data.get("districts", []), "objects": data.get("objects", ""),
        "car_make": data.get("car_make", ""), "car_number": data.get("car_number", ""),
    }
    db.upsert(rec, "fio")
    db.close()