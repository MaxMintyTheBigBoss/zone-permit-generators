# -*- coding: utf-8 -*-
"""
Общий модуль для всех генераторов пропусков (14.3, 14.5, 19.17.1).

Содержит:
- пути (resource_path / runtime_base / template_dir);
- справочник кладбищ и фильтры;
- корректную замену плейсхолдеров с учётом разбивки по run'ам,
  при которой подчёркивается ТОЛЬКО вставленное значение, а метки
  («Фамилия:», «Цель въезда:» и т.п.) остаются обычными;
- вспомогательные join/sanitize/дата.
"""

import os
import re
import sys
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Регулярка для плейсхолдеров: Placeholder_1.1, Placeholder_4.2, **Placeholder_20**
_PH_RE = re.compile(r"\*{0,2}(Placeholder_\d+(?:\.\d+)?)\*{0,2}")


# ---------------------------------------------------------------------------
# Пути
# ---------------------------------------------------------------------------
def resource_path():
    """Путь к ресурсам: совместим с PyInstaller (из .exe берём из _MEIPASS)."""
    base = getattr(sys, "_MEIPASS", None)
    if base:
        return base
    return os.path.dirname(os.path.abspath(__file__))


def runtime_base():
    """Постоянное место рядом с программой (для данных и вывода)."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(os.path.abspath(sys.argv[0]))
    return os.path.dirname(os.path.abspath(__file__))


def template_dir():
    return os.path.join(resource_path(), "templates")


# ---------------------------------------------------------------------------
# Справочник кладбищ
# ---------------------------------------------------------------------------
def load_reference():
    """Справочник: список {district, object}. Объект = 'кладбище о.н.п. <название>'."""
    path = os.path.join(template_dir(), "справочник.docx")
    entries = []
    if not os.path.exists(path):
        return entries
    doc = Document(path)
    for p in doc.paragraphs:
        parts = [t for t in p.text.split("\t") if t.strip()]
        if len(parts) >= 3:
            district = parts[0].strip()
            obj = "кладбище о.н.п. " + parts[2].strip()
            if district and obj:
                entries.append({"district": district, "object": obj})
    return entries


def filter_objects_by_districts(reference, districts):
    """Только кладбища, относящиеся к выбранным районам."""
    allowed = set(districts)
    return [r for r in reference if r["district"] in allowed]


# ---------------------------------------------------------------------------
# Join-помощники
# ---------------------------------------------------------------------------
def join_districts(districts):
    return ", ".join(districts)


def join_objects(objects, custom=""):
    """objects — список записей {object: str} (или строк), custom — произвольный объект."""
    lst = []
    for o in objects:
        if isinstance(o, dict):
            v = (o.get("object") or "").strip()
        else:
            v = (o or "").strip()
        if v:
            lst.append(v)
    if custom and custom.strip():
        lst.append(custom.strip())
    return "; ".join(lst)


def sanitize(s, default="файл"):
    illegal = '<>:"/\\|?*'
    return ("".join(c for c in s if c not in illegal).strip() or default)


def today_dmy():
    return datetime.now().strftime("%d.%m.%Y")


# ---------------------------------------------------------------------------
# Замена плейсхолдеров (исправленная версия)
# ---------------------------------------------------------------------------
def _inline_replace(paragraph, mapping):
    """Заменяет плейсхолдеры в абзаце с учётом разбивки на run'ы.

    Отличия от прежней версии:
    * поиск ведётся по объединённому тексту абзаца, поэтому плейсхолдер,
      разорванный произвольным образом между run'ами, находится корректно;
    * подчёркивается ТОЛЬКО вставленное значение; окружающие метки
      («Фамилия:» и т.п.) сохраняют свой шрифт и остаются обычными;
    * форматирование меток не теряется (вместо «заливания всего текста
      в первый run» абзац перестраивается по сегментам).

    mapping: {token: value}. Плейсхолдер с '*' (например '**Placeholder_20**')
    обрабатывается так же.
    """
    # Собираем run'ы абзаца
    runs = paragraph._p.findall(qn('w:r'))
    if not runs:
        return

    # Полный текст и позиции: список (run, char_range_start_in_run)
    segments = []      # (rpr_element_or_None, text, char_range)
    full_parts = []
    for r in runs:
        texts = [t.text or "" for t in r.findall(qn('w:t'))]
        txt = "".join(texts)
        rpr = r.find(qn('w:rPr'))
        segments.append(_Segment(rpr, txt))
        full_parts.append(txt)
    full = "".join(full_parts)

    if "Placeholder_" not in full:
        return

    # Строим карту позиции -> номер сегмента + смещение
    # Объединяем в один "глобальный" текст с привязкой к сегментам.
    rebuilt = _replace_in_segments(segments, full, mapping)
    if rebuilt is None:
        return

    # Удаляем старые run'ы и вставляем новые из сегментов
    for r in runs:
        paragraph._p.remove(r)
    for seg in rebuilt:
        if not seg.text:
            continue
        run_el = _make_run(seg)
        paragraph._p.append(run_el)


class _Segment:
    __slots__ = ("rpr", "text", "underline")

    def __init__(self, rpr, text, underline=False):
        self.rpr = rpr
        self.text = text
        self.underline = underline


def _replace_in_segments(segments, full, mapping):
    """Разбивает абзац на сегменты текст/значение, находя плейсхолдеры.

    Возвращает список _Segment или None, если ничего не заменилось.
    each сегмент наследует rpr того run'а, где начинается его текст.
    Подчёркивание ставится только на сегменты-значения.
    """
    matches = list(_PH_RE.finditer(full))
    if not matches or not matches[0].group(1) in mapping:
        return None

    # строим опорные позиции сегментов в глобальном тексте
    boundaries = []
    acc = 0
    for s in segments:
        boundaries.append(acc)
        acc += len(s.text)

    def rpr_at(pos):
        # сегмент, содержащий глобальную позицию pos
        idx = 0
        for i, b in enumerate(boundaries):
            if pos >= b:
                idx = i
            else:
                break
        return segments[idx].rpr

    new_segs = []
    pos = 0
    replaced = False
    for m in matches:
        tok = m.group(1)
        if tok not in mapping:
            continue
        val = str(mapping[tok])
        s, e = m.span()
        # текст между прошлым совпадением и текущим
        if s > pos:
            seg_txt = full[pos:s]
            # убираем "голый" пробел если был добавлен регулякой _PH_RE
            if not new_segs or not seg_txt.strip():
                pass
            new_segs.append(_Segment(rpr_at(pos), seg_txt))
        # значение (подчёркнутое)
        new_segs.append(_Segment(rpr_at(s), val, underline=True))
        pos = e
        replaced = True
    if not replaced:
        return None
    if pos < len(full):
        new_segs.append(_Segment(rpr_at(pos), full[pos:]))
    return new_segs


def _make_run(seg):
    """Создаёт новый <w:r>, копируя rPr из исходного сегмента (если был)."""
    rpr_cp = deepcopy(seg.rpr) if seg.rpr is not None else None
    r = OxmlElement('w:r')
    if rpr_cp is not None:
        r.append(rpr_cp)
    if seg.underline:
        if rpr_cp is None:
            rpr_cp = OxmlElement('w:rPr')
            r.append(rpr_cp)
        u = rpr_cp.find(qn('w:u'))
        if u is None:
            u = OxmlElement('w:u')
            rpr_cp.append(u)
        u.set(qn('w:val'), "single")
    t = OxmlElement('w:t')
    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'
    t.set(XML_SPACE, "preserve")
    t.text = seg.text or ""
    r.append(t)
    return r


def fill_doc(doc, mapping):
    """Заменяет плейсхолдеры во всех абзацах и таблицах документа."""
    for para in doc.paragraphs:
        _inline_replace(para, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _inline_replace(para, mapping)
    return doc