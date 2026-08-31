# -*- coding: utf-8 -*-
"""
Генерация документов пропусков по п. 14.5 (вывоз имущества).
Заявление 14.5 + пропуск на вывоз имущества + транспортный пропуск.
Использует общие модули permit_common и permit_db.
"""
import os
import sys
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

from permit_common import (
    resource_path, runtime_base, template_dir,
    load_reference, filter_objects_by_districts,
    join_districts, join_objects, sanitize, today_dmy,
    fill_doc, _PH_RE,
)
import permit_db as dbm

# Цель въезда (Placeholder_6) — фиксированная для 14.5.
GOAL_145 = "для вывоза имущества"

# Районы (Placeholder_4) — те же 8, что и для 14.3
DISTRICTS = [
    "Брагинский", "Буда-Кошелевский", "Ветковский", "Добрушский",
    "Кормянский", "Наровлянский", "Хойникский", "Чечерский",
]

# Кому на подписание (Placeholder_20) — полный список
SIGNERS = [
    "Заместитель начальника отдела Путькова Т.М.",
    "Начальник отдела Соломейчук А.В.",
    "Заместитель начальника главного управления В.О.Шабловский",
    "Главный специалист Гвоздарев А.А.",
    "Главный специалист Геращенко Г.Н.",
    "Главный специалист Колесан А.И.",
    "Главный специалист Курило А.В.",
    "Главный специалист Новик П.Н.",
    "Главный специалист Одиноченко И.В.",
    "Главный специалист Першко А.С.",
]


def _today_dmy():
    return today_dmy()


def make_application_145(data):
    """Заявление 14.5 (вывоз имущества)."""
    doc = Document(os.path.join(template_dir(), "Заявление_14.5.docx"))
    districts = data.get("districts", [])
    m = {
        "Placeholder_1.1": data["last_name"],
        "Placeholder_1.2": data["first_name"],
        "Placeholder_1.3": data["middle_name"],
        "Placeholder_2": data.get("birth_date", ""),
        "Placeholder_3": data.get("id_number", ""),
        "Placeholder_4": join_districts(districts),
        "Placeholder_5": data.get("objects", ""),
        "Placeholder_7": data.get("date_from", ""),
        "Placeholder_8": data.get("date_to", ""),
        "Placeholder_12": data.get("car_make", ""),
        "Placeholder_13": data.get("car_number", ""),
        "Placeholder_21": data.get("cargo", ""),  # вид и количество имущества
    }
    for i in range(3):
        m["Placeholder_4.%d" % (i + 1)] = districts[i] if i < len(districts) else ""

    fill_doc(doc, m)
    return doc


def make_permit_cargo(data):
    """Пропуск на вывоз имущества."""
    doc = Document(os.path.join(template_dir(), "Пропуск_вывоз_имущества.docx"))
    districts = data.get("districts", [])
    m = {
        "Placeholder_1.1": data["last_name"],
        "Placeholder_1.2": data["first_name"],
        "Placeholder_1.3": data["middle_name"],
        "Placeholder_4": join_districts(districts),
        "Placeholder_5": data.get("objects", ""),
        "Placeholder_7": data.get("date_from", ""),
        "Placeholder_8": data.get("date_to", ""),
        "Placeholder_20": data.get("issued_by", ""),
        "Placeholder_21": data.get("cargo", ""),
    }
    fill_doc(doc, m)
    return doc


def make_transport_145(data):
    """Транспортный пропуск для 14.5 (Placeholder_6 всегда "для вывоза имущества")."""
    doc = Document(os.path.join(template_dir(), "Пропуск_транспортный.docx"))
    districts = data.get("districts", [])
    m = {
        "Placeholder_1.1": data["last_name"],
        "Placeholder_1.2": data["first_name"],
        "Placeholder_1.3": data["middle_name"],
        "Placeholder_4": join_districts(districts),
        "Placeholder_5": data.get("objects", ""),
        "Placeholder_6": GOAL_145,
        "Placeholder_7": data.get("date_from", ""),
        "Placeholder_8": data.get("date_to", ""),
        "Placeholder_12": data.get("car_make", ""),
        "Placeholder_13": data.get("car_number", ""),
        "Placeholder_20": data.get("issued_by", ""),
    }
    fill_doc(doc, m)
    return doc


def generate_all_145(data, output_dir=None):
    """Полный комплект документов по п. 14.5. Возвращает список файлов."""
    if output_dir is None:
        output_dir = os.path.join(runtime_base(), "output")
    stamp = datetime.now().strftime("%d.%m.%Y.%H.%M")
    out = os.path.join(output_dir, stamp)
    os.makedirs(out, exist_ok=True)

    created = []
    lname = sanitize(data["last_name"], "Заявитель")

    d = make_application_145(data)
    p = os.path.join(out, "Заявление_14.5_%s.docx" % lname)
    d.save(p)
    created.append(p)

    d = make_permit_cargo(data)
    p = os.path.join(out, "Пропуск_ВывозИмущества_%s.docx" % lname)
    d.save(p)
    created.append(p)

    if data.get("car_make", "").strip() or data.get("car_number", "").strip():
        car = sanitize(data.get("car_number", "") or "Авто", "Авто")
        d = make_transport_145(data)
        p = os.path.join(out, "Пропуск_Транспорт_%s.docx" % car)
        d.save(p)
        created.append(p)

    _db_remember_145(data)
    return created


# ---------------------------------------------------------------------------
# База знаний — SQLite
# ---------------------------------------------------------------------------
def _db_path_145():
    d = os.path.join(runtime_base(), "data")
    os.makedirs(d, exist_ok=True)
    return os.path.join(d, dbm.db_filename("145"))


def _db_conn_145():
    return dbm.PermitDB(_db_path_145())


def _db_load_145():
    try:
        return _db_conn_145().load()
    except Exception:
        return []


def db_find_145(fio):
    return dbm.PermitDB(_db_path_145()).find(fio, "fio") if (fio or "").strip() else None


def _db_remember_145(data):
    db = _db_conn_145()
    try:
        import json as _json
        legacy = os.path.join(runtime_base(), "data", "permit_history_145.json")
        if os.path.exists(legacy):
            recs = _json.load(open(legacy, encoding="utf-8"))
            if recs:
                dbm.migrate_from_json(db, legacy, "fio")
                os.rename(legacy, legacy + ".migrated")
    except Exception:
        pass
    key = " ".join(x for x in [data.get("last_name", ""), data.get("first_name", ""), data.get("middle_name", "")] if x).strip()
    rec = {
        "key": key, "last_name": data.get("last_name", ""), "first_name": data.get("first_name", ""),
        "middle_name": data.get("middle_name", ""), "birth_date": data.get("birth_date", ""),
        "id_number": data.get("id_number", ""), "cargo": data.get("cargo", ""),
        "districts": data.get("districts", []), "objects": data.get("objects", ""),
        "car_make": data.get("car_make", ""), "car_number": data.get("car_number", ""),
    }
    db.upsert(rec, "fio")
    db.close()