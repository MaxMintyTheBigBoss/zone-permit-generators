# -*- coding: utf-8 -*-
"""Тесты для общего модуля замены плейсхолдеров и join-функций."""
import os
import sys
import tempfile

# корневая папка генераторов
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest
from docx import Document
from docx.shared import Pt

import permit_common as pc


class TestInlineReplace:
    def test_simple_placeholder(self):
        """Метка + плейсхолдер в одном run → подчёркнуто только значение."""
        d = Document()
        p = d.add_paragraph("Фамилия: ")
        p.add_run("Placeholder_1.1")
        pc._inline_replace(p, {"Placeholder_1.1": "Иванов"})
        assert p.text == "Фамилия: Иванов"
        assert p.runs[0].text == "Фамилия: "
        assert p.runs[0].font.underline in (None, False)
        assert p.runs[1].text == "Иванов"
        assert p.runs[1].font.underline is True

    def test_split_placeholder(self):
        """Плейсхолдер разорван между run'ами — всё равно заменяется."""
        d = Document()
        p = d.add_paragraph()
        p.add_run("Цель: ")
        p.add_run("Placehol")
        p.add_run("der_6")
        pc._inline_replace(p, {"Placeholder_6": "благоустройство"})
        assert p.text == "Цель: благоустройство"
        # должно быть ровно 2 run'а: метка и значение
        assert len(p.runs) == 2
        assert p.runs[1].font.underline is True

    def test_preserves_label_size(self):
        """Шрифт метки сохраняется, значение получает подчёркивание."""
        d = Document()
        p = d.add_paragraph()
        r1 = p.add_run("Фамилия: ")
        r1.font.size = Pt(14)
        r2 = p.add_run("Placeholder_1.1")
        r2.font.size = Pt(14)
        pc._inline_replace(p, {"Placeholder_1.1": "Иванов"})
        # метка должна сохранить Pt(14)
        assert p.runs[0].font.size == Pt(14)
        assert p.runs[1].font.size == Pt(14)
        assert p.runs[0].font.underline in (None, False)
        assert p.runs[1].font.underline is True

    def test_placeholder_with_stars(self):
        """Плейсхолдер, обёрнутый в **Placeholder_20**, тоже заменяется."""
        d = Document()
        p = d.add_paragraph("Подписал: **Placeholder_20**")
        pc._inline_replace(p, {"Placeholder_20": "Петров"})
        assert "Петров" in p.text
        assert "Placeholder" not in p.text

    def test_no_placeholder_no_change(self):
        d = Document()
        p = d.add_paragraph("Привет, мир")
        before = list(p.runs)
        pc._inline_replace(p, {"Placeholder_1.1": "X"})
        # ничего не должно измениться
        assert p.text == "Привет, мир"
        assert len(p.runs) == len(before)

    def test_empty_mapping_value(self):
        d = Document()
        p = d.add_paragraph("Фамилия: Placeholder_1.1")
        pc._inline_replace(p, {"Placeholder_1.1": ""})
        # должно работать: «Фамилия: » без значения
        assert p.text.startswith("Фамилия: ")


class TestFillDoc:
    def test_fill_in_table(self):
        d = Document()
        tbl = d.add_table(rows=1, cols=2)
        tbl.rows[0].cells[0].text = "Фамилия: Placeholder_1.1"
        tbl.rows[0].cells[1].text = "Имя: Placeholder_1.2"
        pc.fill_doc(d, {"Placeholder_1.1": "Иванов", "Placeholder_1.2": "Иван"})
        assert tbl.rows[0].cells[0].text == "Фамилия: Иванов"
        assert tbl.rows[0].cells[1].text == "Имя: Иван"


class TestJoin:
    def test_join_districts(self):
        assert pc.join_districts(["Брагинский", "Хойникский"]) == "Брагинский, Хойникский"

    def test_join_objects_with_dicts(self):
        objs = [{"object": "A"}, {"object": "B"}]
        assert pc.join_objects(objs) == "A; B"

    def test_join_objects_with_strings(self):
        assert pc.join_objects(["A", "B"]) == "A; B"

    def test_join_objects_with_custom(self):
        assert pc.join_objects([{"object": "A"}], custom="B") == "A; B"

    def test_sanitize(self):
        assert pc.sanitize('a<b>c:d/e\\f|g?h"i*j', "x") == "abcdefghij"


class TestPermitDB:
    def setup_method(self):
        self._tmp = tempfile.mkdtemp()
        self._path = os.path.join(self._tmp, "test.sqlite")
        import permit_db
        self.db = permit_db.PermitDB(self._path)

    def teardown_method(self):
        self.db.close()

    def test_upsert_and_find(self):
        self.db.upsert({"key": "иванов", "last_name": "Иванов"}, "fio")
        rec = self.db.find("Иванов", "fio")  # поиск регистронезависим
        assert rec and rec["last_name"] == "Иванов"

    def test_upsert_updates_existing(self):
        self.db.upsert({"key": "иванов", "last_name": "Иванов", "city": "Минск"}, "fio")
        self.db.upsert({"key": "иванов", "last_name": "Иванов", "city": "Гомель"}, "fio")
        rec = self.db.find("Иванов", "fio")
        assert rec["city"] == "Гомель"

    def test_find_like(self):
        self.db.upsert({"key": "иванов", "last_name": "Иванов"}, "fio")
        self.db.upsert({"key": "петров", "last_name": "Петров"}, "fio")
        matches = self.db.find_like("ива", "fio")
        assert "иванов" in matches
        assert "петров" not in matches

    def test_export_json(self):
        self.db.upsert({"key": "x", "last_name": "X"}, "fio")
        out = os.path.join(self._tmp, "out.json")
        self.db.export(out, "json")
        assert os.path.exists(out)
        import json
        with open(out, encoding="utf-8") as f:
            data = json.load(f)
        assert len(data) == 1 and data[0]["last_name"] == "X"

    def test_export_csv(self):
        self.db.upsert({"key": "x", "last_name": "X"}, "fio")
        out = os.path.join(self._tmp, "out.csv")
        self.db.export(out, "csv")
        assert os.path.exists(out)
        with open(out, encoding="utf-8-sig") as f:
            data = f.read()
        assert "X" in data


class TestGenerators:
    """Интеграционные тесты: генерация всех трёх процедур без ошибок."""

    def setup_method(self):
        self._tmp = tempfile.mkdtemp()
        import permit_db
        # подменяем runtime_base на временную папку через смену frozen/sys
        # (для теста используем явный output_dir)
        self._permit_db = permit_db

    def test_generator_143(self):
        import generator
        data = {
            "last_name": "Тест", "first_name": "Иван", "middle_name": "Иванович",
            "birth_date": "01.01.2000", "id_number": "123",
            "goal": "благоустройство", "districts": ["Брагинский"],
            "objects": "кладбище о.н.п. Тест", "date_from": "01.01.2026",
            "date_to": "31.12.2026", "car_make": "Lada", "car_number": "А001АА7",
            "issued_by": "Иванов", "persons": [],
        }
        files = generator.generate_all(data, self._tmp)
        assert len(files) == 3  # заявление + инд.пропуск + транспорт

    def test_generator_145(self):
        import generator_145 as g
        data = {
            "last_name": "Тест", "first_name": "Иван", "middle_name": "Иванович",
            "birth_date": "01.01.2000", "id_number": "123",
            "cargo": "мебель", "districts": ["Брагинский"],
            "objects": "кладбище о.н.п. Тест", "date_from": "01.01.2026",
            "date_to": "31.12.2026", "car_make": "Lada", "car_number": "А001АА7",
            "issued_by": "Иванов",
        }
        files = g.generate_all_145(data, self._tmp)
        assert len(files) == 3  # заявление + вывоз + транспорт

    def test_generator_19171(self):
        import generator_19171 as g
        data = {
            "org_info": "ООО Тест", "org_short": "Тест",
            "goal": "мониторинг", "districts": ["Брагинский"],
            "objects": "ГПНИУ ПГРЭЗ", "date_from": "01.01.2026",
            "date_to": "31.12.2026", "issued_by": "Иванов",
            "persons": [{"last_name": "Иванов", "first_name": "Иван",
                         "middle_name": "Иванович", "position": "инженер"}],
            "vehicles": [{"make": "Lada", "number": "А001АА7"}],
        }
        files = g.generate_all_19171(data, self._tmp)
        assert len(files) == 2  # 1 инд. + 1 транспорт